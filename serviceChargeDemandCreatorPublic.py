from docx import Document
import os
from docx2pdf import convert
import time
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from datetime import datetime
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from openpyxl import load_workbook
import smtplib
import mimetypes
from email.message import EmailMessage

# Load the Excel file
excel_file = r"data.xlsx"
wb = load_workbook(excel_file)
sheet = wb.active  # Use the first sheet

for row in sheet.iter_rows(min_row=2, values_only=True):
    if not row[0]:  # Skip empty rows in the first column
        continue

    #Set variables from the Excel file
    flat_no= f"{row[0]}" # Use first column value as filename
    ownerName = f"{row[1]}"  # Use second column value as owner name
    arrears = float(f"{row[2]}")  # Use third column value as arrears
    service_charge = float(f"{row[3]}")  # Use fourth column value as service charge
    first_name = f"{row[4]}" # Get the first name of the owner
    email = str(f"{row[5]}") # Get the email address of the owner 
    document = Document()  # Create a new Word document

    # Get the current month and year
    today = datetime.today()
    current_month = today.month
    current_year = today.year

    # Determine the next April (4) or October (10), whichever is sooner
    if current_month < 4:
        next_month = 4  # April is next
    elif current_month < 10:
        next_month = 10  # October is next
    else:
        next_month = 4  # April next year
        current_year += 1  # Move to the next year

    # Format the date as "Month Year" (e.g., "Apr 2024")
    next_month_name = datetime(current_year, next_month, 1).strftime("%b %y")
    till_month = next_month + 6
    till_year = today.year
    if till_month == 4 :
        till_year + 1
    till_month_name = datetime( till_year, till_month, 1).strftime("%b %y")
    month_year = datetime(current_year,next_month,1).strftime("%B %Y")

    # Set font
    style = document.styles['Normal']
    style.font.name = 'Calibri'

    # Add Header with two columns 
    header_section = document.sections[0]
    header_section.different_first_page_header_footer = True
    header = header_section.first_page_header
    # create header table
    table = header.add_table(rows=1, cols=2, width=Cm(30))
    table.columns[0].width = Cm(9)
    table.columns[1].width = Cm(9)
    #left header column
    left_cell = table.cell(0, 0)
    left_paragraph = left_cell.paragraphs[0]
    run = left_paragraph.add_run(f'{ownerName}\n') 
    run.bold = True
    left_paragraph.add_run(f'{flat_no} The Lodge\nWoody Road\nLondon\nW1 7PR')
    left_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    #Right Header column
    right_cell = table.cell(0, 1)
    right_paragraph = right_cell.paragraphs[0]
    run = right_paragraph.add_run(f'\n\nIssue Date: {datetime.now().strftime("%d %B %Y")}\n')
    run.bold = True
    right_paragraph.add_run(f'Invoice Reference: S/C {next_month_name} Flat {flat_no}')
    right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    #footer expermint
    footer = header_section.first_page_footer
    footer_para = footer.paragraphs[0]
    footer_para.text = 'Pursuant to the Landlord and Tenant Act 1987 sections 47 and 48 that your Landlord’s name and registered address is THE LODGE LIMITED, C/O Company Secretary Ltd, The Lodge, Woody Road, London, W1 7PR England, registration number No: 1234560 and all notices (including notice of proceedings) should be served upon the landlord at their stated registered office address: THE LODGE LIMITED, C/O Company Secretary Ltd, The Lodge, Woody Road, London, W1 7PR England'
    run = footer_para.runs[0]
    run.font.size = Pt(8)
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # first Paragraphs
    firstParagraph = document.add_paragraph()
    firstParagraph.add_run('Invoice for Service Charge Payment\n').bold = True
    firstParagraph.add_run(f'Re: Flat {flat_no}, The Lodge, Woody Road, London, W1 7PR\n\n\n')
    firstParagraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # add breakdown table
    mainTable = document.add_table(rows=4, cols=2)
    for row in mainTable.rows:
        row.cells[0].width = Cm(25)
        row.cells[1].width = Cm(3)

    tl_cell = mainTable.cell(0, 0)  # tl_cell = top left cell
    tl_cell_paragraph = tl_cell.paragraphs[0]
    tl_cell_paragraph.add_run('Description').bold = True
    sl_cell = mainTable.cell(1, 0)  # sl = 2nd left cell
    sl_cell_paragraph = sl_cell.paragraphs[0]
    sl_cell_paragraph.add_run('Balance brought forward')
    tl_cell = mainTable.cell(2, 0)
    tl_cell_paragraph = tl_cell.paragraphs[0]
    tl_cell_paragraph.add_run(f'Service charge for the period starting on {next_month_name} to {till_month_name}')
    bl_cell = mainTable.cell(3, 0)  # bl cell = bottom left cell
    bl_cell_paragraph = bl_cell.paragraphs[0]
    bl_cell_paragraph.add_run('Total Amount Due')
    tr_cell = mainTable.cell(0, 1)  # tr_cell - top right cell
    tr_cell_paragraph = tr_cell.paragraphs[0]
    tr_cell_paragraph.add_run('Value').bold = True
    sr_cell = mainTable.cell(1, 1)  # sr_cell = 2nd right cell
    sr_cell_paragraph = sr_cell.paragraphs[0]
    sr_cell_paragraph.add_run(f'£{arrears:.2f}')
    tr_cell = mainTable.cell(2, 1)  # tr_cell = 3rd right cell
    tr_cell_paragraph = tr_cell.paragraphs[0]
    tr_cell_paragraph.add_run(f'£{service_charge:.2f}')
    lr_cell = mainTable.cell(3, 1)  # lr_cell = last right cell
    lr_cell_paragraph = lr_cell.paragraphs[0]
    lr_cell_paragraph.add_run(f'£{arrears + service_charge:.2f}').underline = True

    # Add borders to the table
    tbl = mainTable._element
    tblBorders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        r'<w:insideH w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        r'<w:insideV w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        r'</w:tblBorders>'
    )
    tbl.tblPr.append(tblBorders)

    # Reduce spacing after each paragraph in the table
    for row in mainTable.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)

    # payment paragraphs
    payPar = document.add_paragraph('\n\n')
    run = payPar.add_run(' Payment may be made to:\n\n').bold = True
    run = payPar.add_run('WOODY HOUSE FREEHOLD LIMITED\n\n')
    run = payPar.add_run(' Bank Details:\n\n').bold = True
    run = payPar.add_run(f'Account Name: WOODY HOUSE FREEHOLD LIMITED\nAccount Number: 1234567\nSort Code: 04-07-09')

    # 2nd page
    document.add_page_break()

    raoText = """ 1. This summary, which briefly sets out your rights and obligations in relation to    ariable service charges, must by law accompany a demand for service charges. Unless a summary is sent to you with a demand, you may withhold the service charge. The summary does not give a full interpretation of the law and if you are in doubt about your rights and obligations you should seek independent advice.
    2. Your lease sets out your obligations to pay service charges to your landlord in addition to your rent. Service charges are amounts payable for services, repairs, maintenance, improvements, insurance or the landlord’s costs of management, to the extent that the costs have been reasonably incurred.
    3.You have the right to ask a “First-tier Tribunal (Property Chamber)” to determine whether you are liable to pay service charges for services, repairs, maintenance, improvements, insurance or management. You may make a request before or after you have paid the service charge. If the tribunal determines that the service charge is payable, the tribunal may also determine: Who should pay the service charge and who it would be paid to;
    The amount; The date it should be paid by; and how it should be paid. However, you do not have these rights where:
    A matter has been agreed or admitted by you;
    A matter has already been, or is to be referred to arbitration or has been determined by arbitration and you agreed to go to arbitration after the disagreement about the service charge or costs arose;
    Or a matter has been decided by a court.
    4. If your lease allows your landlord to recover costs incurred or that may be incurred in legal proceedings as service charges, you may ask the court or tribunal, before which those proceedings were brought, to rule that your landlord may not do so.
    5. Where you seek a determination from a “First-tier Tribunal (Property Chamber)”, you will have to pay an application fee and, where the matter proceeds to a hearing, a hearing fee, unless you qualify for a waiver or reduction. The total fees payable will not exceed £500, but making an application may incur additional costs, such as professional fees, which you may also have to pay.
    6. A “First-tier Tribunal (Property Chamber)” has the power to award costs, not exceeding £500, against a party to any proceedings where –
    It dismisses a matter because it is frivolous, vexatious or an abuse of process; or
    It considers a party has acted frivolously, vexatiously, abusively, disruptively or unreasonably.
    The Upper Tribunal has similar powers when hearing an appeal against a decision of a “First-tier Tribunal (Property Chamber)”.
    7. If your landlord – Proposes works on a building or any other premises that will cost you or any other tenant more than £250, or proposes to enter into an agreement for works or services which will last for more than 12 months and will cost you or any other tenant more than £100 in any 12 month accounting period.
    Your contribution will be limited to these amounts unless your landlord has properly consulted on the proposed works or agreement or a “First-tier Tribunal (Property Chamber)” has agreed that consultation is not required.
    8. You have the right to apply to a “First-tier Tribunal (Property Chamber)” to ask it to determine whether your lease should be varied on the grounds that it does not make satisfactory provision in respect of the calculation of a service charge payable under the lease.
    9. You have the right to write to your landlord to request a written summary of the costs which make up the service charges. The summary must –
    Cover the last 12 month period used for making up the accounts relating to the service charge ending no later than the date of your request, where the accounts are made up for 12 month periods; or
    Cover the 12 month period ending with the date of your request, where the accounts are not made up for 12 month periods. The summary must be given to you within 1 month of your request or 6 months of the end of the period to which the summary relates whichever is the later.
    10. You have the right, within 6 months of receiving a written summary of costs, to require the landlord to provide you with reasonable facilities to inspect the accounts, receipts and other documents supporting the summary and for taking copies or extracts from them.
    11. You have the right to ask an accountant or surveyor to carry out an audit of the financial management of the premises containing your dwelling, to establish the obligations of your landlord and the extent to which the service charges you pay are being used efficiently. It will depend on your circumstances whether you can exercise this right alone or only with the support of others living in the premises. You are strongly advised to seek independent advice before exercising this right.
    12. Your lease may give your landlord a right of re-entry or forfeiture where you have failed to pay charges which are properly due under the lease. However, to exercise this right, the landlord must meet all the legal requirements and obtain a court order. A court order will only be granted if you have admitted you are liable to pay the amount or it is finally determined by a Court, tribunal or by arbitration that the amount is due. The court has a wide discretion in granting such an order and it will take into account all the circumstances of the case.
    """

    #Add rights and obligations text
    raoPar = document.add_paragraph()
    run = raoPar.add_run("Service Charges—Summary of Tenants’ Rights and Obligations").bold = True
    raoPar.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    raoPar2 = document.add_paragraph()
    run = raoPar2.add_run(raoText)
    run.font.size = Pt(8)

    #Save to Word Documentnd convert to PDF
    file_name = (f'Flat {flat_no} {month_year}')
    word_fn = f"{file_name}" + '.docx'
    pdf_fn = f"{file_name}" + '.pdf'
    document.save(word_fn)
    time.sleep (2)
    convert(word_fn, pdf_fn)
    os.startfile(pdf_fn)

        # ✅ Function to send an email with a PDF attachment
    def send_email_with_attachment(email, subject, body, pdf_fn):
        GMAIL_USER = "directors@ghfl.com"
        GMAIL_PASSWORD = "your-app-password-here"  # ✅ Use Gmail App Password

        msg = EmailMessage()
        msg["From"] = GMAIL_USER
        msg["To"] = email
        msg["Subject"] = subject
        msg.set_content(body)

        # ✅ Get MIME type for the PDF file
        mime_type, encoding = mimetypes.guess_type(pdf_fn)
        if mime_type is None:
            mime_type = "application/pdf"  # Default to PDF type
        main_type, sub_type = mime_type.split("/", 1)

        # ✅ Attach PDF file
        try:
            with open(pdf_fn, "rb") as file:
                msg.add_attachment(file.read(), maintype=main_type, subtype=sub_type, filename=os.path.basename(pdf_fn))
        except FileNotFoundError:
            print(f"🚨 Error: File '{pdf_fn}' not found. Email not sent.")
            return  # Stop execution for this email

        # ✅ Send Email via Gmail SMTP
        try:
            with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
                server.login(GMAIL_USER, GMAIL_PASSWORD)
                server.send_message(msg)
            print(f"✅ Email sent to {email} with attachment {pdf_fn}")
        except Exception as e:
            print(f"🚨 Email sending failed: {e}")


    # ✅ Send Email
    send_email_with_attachment(
        email,
        subject=f"Invoice - {file_name}",
        body=f"Dear {first_name},\n\nPlease find attached your invoice for {next_month_name}.\n\nBest Regards,\nGenesis House Freehold Ltd.",
        pdf_fn=pdf_fn
    )

    # Print a message to the console
    print('Word and pdf documents created and emails sent!')

